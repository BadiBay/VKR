from rest_framework import viewsets, status
from rest_framework.decorators import action
from rest_framework.response import Response
from django.http import JsonResponse, HttpResponse
from django.core.cache import cache
from django.db.models import Sum
from celery.result import AsyncResult
import requests
from bs4 import BeautifulSoup
import re
import statistics
import time
import urllib.parse
from docx import Document 
from docx.shared import Pt 
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
from urllib.parse import urlparse, parse_qs, unquote

from gigachat import GigaChat
from sklearn.feature_extraction.text import TfidfVectorizer
import numpy as np

from .models import Project, Keyword, Cluster, AuditLog, APIKey, APILog, AIRole
from .serializers import (ProjectSerializer, ProjectDetailSerializer, 
                          ClusterSerializer, AuditLogSerializer, 
                          APIKeySerializer, APILogSerializer, AIRoleSerializer)
from .tasks import fetch_keywords_for_project, cluster_keywords_for_project

def get_gigachat_key():
    key = APIKey.objects.filter(is_active=True, name__icontains='gigachat').first()
    if key: return key.key
    # Fallback to the one in code
    return "MDE5YWI3NTgtMWE0MC03NGU3LWI4YzgtMjM2NDJmOGM2M2ZjOjY5NWE2MjkzLWU4ZDgtNGM0OC1hNmZhLTg3YmZiZjRlNWY5OA=="

def log_api_call(request, endpoint, duration, status_msg, tokens=0):
    ip = request.META.get('REMOTE_ADDR')
    APILog.objects.create(user_ip=ip, endpoint=endpoint, duration_ms=duration, status=status_msg, tokens_used=tokens)

def search_engine_parser(query, limit=50):
    print(f"🔎 Позиции: ищем '{query}'...")
    urls = []
    try:
        google_url = "https://www.google.com/search"
        params = {'q': query, 'num': limit, 'hl': 'ru'}
        headers = {
            'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8'
        }
        resp = requests.get(google_url, params=params, headers=headers, timeout=5)
        if resp.status_code == 200:
            soup = BeautifulSoup(resp.text, 'html.parser')
            for a in soup.find_all('a'):
                href = a.get('href')
                if href and href.startswith('http') and 'google' not in href:
                    urls.append(href)
            if len(urls) > 5:
                print(f"✅ Google выдал {len(urls)} результатов.")
                return urls[:limit]
    except Exception as e:
        print(f"⚠️ Google search error: {e}")

    print("🔄 Переключаемся на DuckDuckGo...")
    try:
        ddg_url = "https://html.duckduckgo.com/html/"
        headers = {'User-Agent': 'Mozilla/5.0'}
        resp = requests.post(ddg_url, data={'q': query}, headers=headers, timeout=10)
        soup = BeautifulSoup(resp.text, 'html.parser')
        for link in soup.find_all('a', class_='result__a'):
            href = link.get('href')
            if href:
                if 'uddg=' in href:
                    parsed = urlparse(href)
                    qs = parse_qs(parsed.query)
                    if 'uddg' in qs:
                        urls.append(qs['uddg'][0])
                else:
                    urls.append(href)
        print(f"✅ DDG выдал {len(urls)} результатов.")
    except Exception as e:
        print(f"⚠️ DDG search error: {e}")

    return urls[:limit]

def analyze_serp(query):
    print(f"Анализируем конкурентов для: {query}")
    headers = {'User-Agent': 'Mozilla/5.0'}
    STOP_DOMAINS = ['ozon.ru', 'wildberries.ru', 'market.yandex.ru', 'avito.ru', 'youtube.com', 'vk.com', 'wikipedia.org']
    raw_urls = search_engine_parser(query, limit=30)
    competitor_urls = []
    for url in raw_urls:
        if any(bad in url.lower() for bad in STOP_DOMAINS): continue
        competitor_urls.append(url)
        if len(competitor_urls) >= 3: break

    text_lengths, all_headers, page_texts = [], [], []
    for url in competitor_urls:
        try:
            resp = requests.get(url, headers=headers, timeout=5)
            soup = BeautifulSoup(resp.text, 'html.parser')
            for s in soup(["script", "style", "nav", "footer", "header", "aside"]): s.extract()
            text = soup.get_text(separator=' ')
            clean_text = re.sub(r'\s+', ' ', text).strip()
            if len(clean_text) > 500:
                text_lengths.append(len(clean_text))
                page_texts.append(clean_text)
                headers_list = []
                for h in soup.find_all(['h2', 'h3']):
                    txt = h.get_text().strip()
                    if 5 < len(txt) < 100: headers_list.append(f"{h.name.upper()}: {txt}")
                if headers_list:
                    domain = url.split('//')[-1].split('/')[0]
                    all_headers.append(f"\n>>> {domain}:\n" + "\n".join(headers_list[:10]))
        except: continue

    avg_len = int(statistics.mean(text_lengths)) if text_lengths else 3500
    lsi_words = []
    if page_texts:
        try:
            vectorizer = TfidfVectorizer(max_features=20, token_pattern=r'(?u)\b[а-яА-ЯёЁ]{4,}\b')
            tfidf_matrix = vectorizer.fit_transform(page_texts)
            feature_names = vectorizer.get_feature_names_out()
            sums = tfidf_matrix.sum(axis=0)
            data = [(term, sums[0, col]) for col, term in enumerate(feature_names)]
            lsi_words = [x[0] for x in sorted(data, key=lambda x: x[1], reverse=True)[:15]]
        except Exception as e: print(f"TF-IDF Error: {e}")

    return {
        'competitors_urls': competitor_urls,
        'avg_text_length': avg_len,
        'competitors_structure': "\n".join(all_headers),
        'lsi_words': lsi_words
    }

def check_site_health(url):
    print(f"Аудит сайта: {url}")
    if not url.startswith('http'): url = 'https://' + url
    stats = {'score': 100, 'checks': []}
    def add_check(name, status, msg, penalty=0):
        stats['checks'].append({'name': name, 'status': status, 'msg': msg})
        if not status: stats['score'] = max(0, stats['score'] - penalty)

    try:
        start_time = time.time()
        headers = {'User-Agent': 'Mozilla/5.0'}
        resp = requests.get(url, headers=headers, timeout=10)
        load_time = round((time.time() - start_time) * 1000)
        
        if resp.status_code == 200: add_check('Доступность', True, '200 OK')
        else: add_check('Доступность', False, f'Код {resp.status_code}', 50)

        if url.startswith('https'): add_check('SSL', True, 'HTTPS есть')
        else: add_check('SSL', False, 'Нет HTTPS', 20)
            
        if load_time < 500: add_check('Скорость', True, f'{load_time} мс')
        elif load_time < 1500: add_check('Скорость', True, f'{load_time} мс (Норм)')
        else: add_check('Скорость', False, f'{load_time} мс (Медленно)', 10)

        soup = BeautifulSoup(resp.text, 'html.parser')
        
        if soup.title and soup.title.string:
            l = len(soup.title.string)
            if 10 < l < 70: add_check('Title', True, f'OK ({l} симв.)')
            else: add_check('Title', False, f'Неоптимально ({l} симв.)', 5)
        else: add_check('Title', False, 'Отсутствует!', 20)

        desc = soup.find('meta', attrs={'name': 'description'})
        if desc and desc.get('content'): add_check('Description', True, 'Присутствует')
        else: add_check('Description', False, 'Отсутствует!', 10)

        h1s = soup.find_all('h1')
        if len(h1s) == 1: add_check('H1', True, 'Один H1 (Отлично)')
        elif len(h1s) == 0: add_check('H1', False, 'Нет H1', 15)
        else: add_check('H1', False, f'Найдено {len(h1s)} H1', 10)
    except Exception as e:
        add_check('Ошибка', False, str(e)[:50], 100)
        stats['score'] = 0

    return stats

class ClusterViewSet(viewsets.ModelViewSet):
    queryset = Cluster.objects.all()
    serializer_class = ClusterSerializer

    @action(detail=True, methods=['post'])
    def change_status(self, request, pk=None):
        cluster = self.get_object()
        new_status = request.data.get('status')
        if new_status in dict(Cluster.STATUS_CHOICES):
            cluster.status = new_status
            cluster.save()
            return Response({'status': cluster.status})
        return Response({'error': 'Invalid status'}, status=400)

    @action(detail=True, methods=['post'])
    def move_keyword(self, request, pk=None):
        cluster = self.get_object()
        keyword_id = request.data.get('keyword_id')
        try:
            kw = Keyword.objects.get(id=keyword_id)
            kw.cluster = cluster
            kw.save()
            return Response({'status': 'Keyword moved'})
        except Keyword.DoesNotExist:
            return Response({'error': 'Keyword not found'}, status=404)

class APIKeyViewSet(viewsets.ModelViewSet):
    queryset = APIKey.objects.all()
    serializer_class = APIKeySerializer

class APILogViewSet(viewsets.ModelViewSet):
    queryset = APILog.objects.all().order_by('-created_at')
    serializer_class = APILogSerializer

class AIRoleViewSet(viewsets.ModelViewSet):
    queryset = AIRole.objects.all()
    serializer_class = AIRoleSerializer

class ProjectViewSet(viewsets.ModelViewSet):
    queryset = Project.objects.all().order_by('-created_at')

    # ------------------------------------------------------------------ #
    # Ключи кэша                                                           #
    # ------------------------------------------------------------------ #
    CACHE_KEY_LIST = 'project_list'
    CACHE_TIMEOUT = 900  # 15 минут

    @staticmethod
    def _project_detail_key(pk):
        return f'project_detail_{pk}'

    def _invalidate_project_caches(self, pk=None):
        """Удаляет кэш списка проектов и, при наличии pk, детальный кэш."""
        cache.delete(self.CACHE_KEY_LIST)
        if pk is not None:
            cache.delete(self._project_detail_key(pk))

    def get_serializer_class(self):
        if self.action == 'retrieve': return ProjectDetailSerializer
        return ProjectSerializer

    # ------------------------------------------------------------------ #
    # Переопределяем стандартные CRUD-методы для кэширования              #
    # ------------------------------------------------------------------ #
    def list(self, request, *args, **kwargs):
        """
        GET /api/projects/ — список проектов с кэшированием в Redis.
        Первый запрос → достаём из PostgreSQL, сохраняем в кэше.
        Повторные запросы → мгновенный ответ из Redis (≈1 мс вместо ~50 мс).
        """
        cached = cache.get(self.CACHE_KEY_LIST)
        if cached is not None:
            return Response(cached)
        response = super().list(request, *args, **kwargs)
        cache.set(self.CACHE_KEY_LIST, response.data, self.CACHE_TIMEOUT)
        return response

    def retrieve(self, request, *args, **kwargs):
        """
        GET /api/projects/<id>/ — детальный вид проекта (с кластерами / ключевиками).
        Кэшируем раздельно для каждого проекта.
        """
        pk = kwargs.get('pk')
        cache_key = self._project_detail_key(pk)
        cached = cache.get(cache_key)
        if cached is not None:
            return Response(cached)
        response = super().retrieve(request, *args, **kwargs)
        cache.set(cache_key, response.data, self.CACHE_TIMEOUT)
        return response

    def perform_create(self, serializer):
        """После создания проекта сбрасываем кэш списка."""
        super().perform_create(serializer)
        cache.delete(self.CACHE_KEY_LIST)

    def perform_update(self, serializer):
        """После обновления сбрасываем и список, и детальный кэш проекта."""
        instance = serializer.save()
        self._invalidate_project_caches(pk=instance.pk)

    def perform_destroy(self, instance):
        """Перед удалением сбрасываем кэши."""
        self._invalidate_project_caches(pk=instance.pk)
        super().perform_destroy(instance)

    @action(detail=True, methods=['post'])
    def uncategorize_keyword(self, request, pk=None):
        project = self.get_object()
        keyword_id = request.data.get('keyword_id')
        try:
            kw = Keyword.objects.get(id=keyword_id, project=project)
            kw.cluster = None
            kw.save()
            return Response({'status': 'Keyword uncategorized'})
        except Keyword.DoesNotExist:
            return Response({'error': 'Keyword not found'}, status=404)

    @action(detail=True, methods=['post'])
    def import_keywords(self, request, pk=None):
        project = self.get_object()
        file_obj = request.FILES.get('file')
        if not file_obj:
            return Response({'error': 'No file provided'}, status=400)
        
        try:
            if file_obj.name.endswith('.csv'):
                df = pd.read_csv(file_obj)
            else:
                df = pd.read_excel(file_obj)
            
            # Предположим колонки 'keyword' (или 'query') и 'volume'
            query_col = next((c for c in df.columns if 'query' in c.lower() or 'keyword' in c.lower() or 'слово' in c.lower() or 'фраза' in c.lower()), None)
            vol_col = next((c for c in df.columns if 'vol' in c.lower() or 'частот' in c.lower() or 'ws' in c.lower()), None)
            
            if not query_col:
                return Response({'error': 'Could not find keyword column'}, status=400)
                
            keywords = df[query_col].dropna().astype(str).tolist()
            volumes = df[vol_col].fillna(0).astype(int).tolist() if vol_col else [0]*len(keywords)

            # Constraint: 15,000 max
            if Keyword.objects.filter(project=project).count() + len(keywords) > 15000:
                return Response({'error': 'Limit of 15000 keywords per project exceeded'}, status=400)
            
            kws = []
            for q, v in zip(keywords, volumes):
                q = q.strip()[:250]
                if q: kws.append(Keyword(project=project, query=q, volume=v))
            
            Keyword.objects.bulk_create(kws, ignore_conflicts=True)
            return Response({'status': f'Imported {len(kws)} keywords'})
        except Exception as e:
            return Response({'error': str(e)}, status=400)

    @action(detail=True, methods=['post'])
    def mass_delete(self, request, pk=None):
        project = self.get_object()
        filters = request.data.get('filters', {})
        min_volume = filters.get('min_volume')
        stop_words = filters.get('stop_words', [])
        
        qs = project.keywords.all()
        deleted = 0
        
        if min_volume is not None:
            res = qs.filter(volume__lt=int(min_volume)).delete()
            deleted += res[0]
            
        if stop_words:
            for word in stop_words:
                res = qs.filter(query__icontains=word).delete()
                deleted += res[0]

        return Response({'status': f'Deleted {deleted} keywords'})

    @action(detail=True, methods=['post'], url_path='reset_clusters')
    def reset_clusters(self, request, pk=None):
        project = self.get_object()
        only_drafts = request.data.get('only_drafts', False)
        if only_drafts:
            from .models import Cluster
            Cluster.objects.filter(project=project, status='draft').delete()
            return Response({'status': 'Удалены только черновики кластеров'})
        else:
            from .models import Cluster
            Cluster.objects.filter(project=project).delete()
            return Response({'status': 'Все кластеры удалены'})

    @action(detail=True, methods=['post'], url_path='clear_project')
    def clear_project(self, request, pk=None):
        project = self.get_object()
        from .models import Keyword, Cluster
        Keyword.objects.filter(project=project).delete()
        Cluster.objects.filter(project=project).delete()
        # Данные проекта изменились — сбрасываем его кэш
        self._invalidate_project_caches(pk=project.pk)
        return Response({'status': 'Проект очищен от слов и кластеров'})

    @action(detail=True, methods=['post'])
    def run_fetch(self, request, pk=None):
        fetch_keywords_for_project.delay(self.get_object().id)
        return Response({'status': 'started'})

    @action(detail=True, methods=['post'])
    def run_clustering(self, request, pk=None):
        task = cluster_keywords_for_project.delay(self.get_object().id)
        return Response({'task_id': task.id})

    @action(detail=False, methods=['post'])
    def analyze_competitors(self, request):
        query = request.data.get('query', '')
        if not query: return Response({'error': 'No query'}, status=400)
        return Response(analyze_serp(query))

    @action(detail=True, methods=['get'])
    def audit(self, request, pk=None):
        project = self.get_object()
        if not project.url: return Response({'error': 'No URL'}, status=400)
        result = check_site_health(project.url)
        # Сохраняем аудит (PT-3.1)
        AuditLog.objects.create(project=project, score=result['score'], checks_json=result['checks'])
        return Response(result)

    @action(detail=True, methods=['post'])
    def check_position(self, request, pk=None):
        project = self.get_object()
        query = request.data.get('query')
        if not query or not project.url: return Response({'error': 'No data'}, status=400)
        target_domain = project.url.lower().replace('https://', '').replace('http://', '').replace('www.', '').split('/')[0]
        urls = search_engine_parser(query, limit=50)
        position, found_url = -1, ""
        for i, url in enumerate(urls):
            clean_url = unquote(url.lower())
            if target_domain in clean_url:
                position, found_url = i + 1, url
                break
        return Response({'position': str(position) if position > 0 else "> 50", 'found_url': found_url})

    @action(detail=True, methods=['post'])
    def download_docx(self, request, pk=None):
        project = self.get_object()
        cluster_id = request.data.get('cluster_id')
        doc_type = request.data.get('doc_type', 'tz') # 'tz' or 'article'
        text_content = request.data.get('content', '')
        meta_content = request.data.get('meta', '')
        lsi_content = request.data.get('lsi', '')

        if not cluster_id: return Response({'error': 'No cluster_id'}, status=400)
        
        try:
            cluster = Cluster.objects.get(id=cluster_id)
        except Cluster.DoesNotExist:
            return Response({'error': 'Cluster empty or not found'}, status=404)
            
        main_keyword = cluster.keywords.first().query if cluster.keywords.exists() else 'Без названия'
        cluster_name = main_keyword.capitalize()

        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)

        title_text = f'ТЗ для копирайтера: {cluster_name}' if doc_type == 'tz' else f'Статья: {cluster_name}'
        head = doc.add_heading(title_text, 0)
        head.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if meta_content:
            doc.add_heading('SEO-теги (Meta)', level=1)
            doc.add_paragraph(meta_content)

        if doc_type == 'article':
            doc.add_heading('Контент', level=1)
            # Если это статья (aiContent), просто пишем ее по абзацам или как один большой текст
            lines = text_content.split('\n')
            for line in lines:
                if line.strip():
                    doc.add_paragraph(line.strip())
        else:
            # Если это ТЗ (finalText), записываем как есть
            lines = text_content.split('\n')
            for line in lines:
                if line.strip():
                    doc.add_paragraph(line.strip())

        if lsi_content:
            doc.add_heading('LSI слова', level=1)
            doc.add_paragraph(lsi_content)

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        prefix = "Article" if doc_type == 'article' else "TZ"
        filename = f"{prefix}_{cluster_name}.docx"
        response['Content-Disposition'] = f'attachment; filename*=UTF-8\'\'{urllib.parse.quote(filename)}'
        doc.save(response)
        return response

    @action(detail=True, methods=['post'])
    def generate_content(self, request, pk=None):
        project = self.get_object()
        cluster_id = request.data.get('cluster_id')
        role_id = request.data.get('role_id')
        custom_lsi = request.data.get('custom_lsi', '')

        if not cluster_id: return Response({'error': 'No cluster_id'}, status=400)
        
        try:
            cluster = Cluster.objects.get(id=cluster_id)
        except Cluster.DoesNotExist:
            return Response({'error': 'Empty or Missing Cluster'}, status=404)
        
        keywords = cluster.keywords.all().order_by('-volume')
        if not keywords.exists(): return Response({'error': 'Empty keywords'}, status=404)
            
        main_keyword = keywords.first().query
        cluster_name = main_keyword.capitalize()
        top_keys = [k.query for k in keywords[:5]]
        if custom_lsi: top_keys.extend([x.strip() for x in custom_lsi.split(',')])

        role_prompt = "Ты ведущий эксперт и редактор."
        if role_id:
            try:
                role = AIRole.objects.get(id=role_id)
                role_prompt = role.prompt_addition
            except AIRole.DoesNotExist: pass

        t0 = time.time()
        try:
            chat = GigaChat(credentials=get_gigachat_key(), verify_ssl_certs=False, model='GigaChat:latest', temperature=0.8)
            prompt = (
                f"РОЛЬ: {role_prompt}. Твоя задача — написать информативную, экспертную и полезную SEO-статью.\n"
                f"ОСНОВНАЯ ТЕМА: {cluster_name}\n"
                f"ОБЯЗАТЕЛЬНЫЕ КЛЮЧЕВЫЕ СЛОВА (используй органично, можно склонять): {', '.join(top_keys)}\n\n"
                f"ТРЕБОВАНИЯ К ТЕКСТУ:\n"
                f"1. Объем: не менее 8000 символов. Тема должна быть раскрыта глубоко и подробно.\n"
                f"2. Структура: используй структуру статьи с введением, основным телом (заголовки H2, H3) и логичным заключением.\n"
                f"3. Оформление: разделяй текст на удобные абзацы (не более 4-5 предложений), обязательно используй маркированные списки для перечислений.\n"
                f"4. Стиль и польза: пиши естественным, понятным языком для людей. Избегай 'воды', сложных штампов и канцеляризмов. Давай конкретную пользу читателю.\n"
                f"Сгенерируй готовую статью в формате Markdown."
            )
            response = chat.chat(prompt)
            content = response.choices[0].message.content
            
            duration = int((time.time() - t0) * 1000)
            log_api_call(request, "gigachat_content", duration, "200 OK", 500) # mockup tokens
            
            return Response({'content': content})
        except Exception as e:
            log_api_call(request, "gigachat_content", 0, f"Error: {str(e)[:40]}", 0)
            return Response({'content': f"Ошибка генерации: {str(e)}"}, status=500)

    @action(detail=True, methods=['post'])
    def generate_meta(self, request, pk=None):
        main_keyword = request.data.get('keyword')
        role_id = request.data.get('role_id')
        role_prompt = "Ты SEO-специалист."
        if role_id:
            try:
                role = AIRole.objects.get(id=role_id)
                role_prompt = f"Действуй как {role.name}. {role.prompt_addition}"
            except: pass

        t0 = time.time()
        try:
            chat = GigaChat(credentials=get_gigachat_key(), verify_ssl_certs=False, model='GigaChat:latest')
            prompt = (
                f"РОЛЬ: {role_prompt}. Твоя задача — составить качественные, кликабельные SEO метатеги.\n"
                f"ГЛАВНЫЙ ЗАПРОС СТРАНИЦЫ: '{main_keyword}'\n\n"
                f"Сгенерируй 3 уникальных варианта связки (Title + Description).\n"
                f"ТРЕБОВАНИЯ:\n"
                f"1. Title (заголовок): длина от 40 до 60 символов. Должен естественно включать главный запрос ближе к началу и привлекать внимание.\n"
                f"2. Description (описание): длина от 130 до 160 символов. Должен раскрывать суть страницы, содержать призыв к действию (CTA) и отражать ценность для пользователя.\n"
                f"Верни ответ в удобном читабельном текстовом формате без лишних вступлений."
            )
            res = chat.chat(prompt)
            duration = int((time.time() - t0) * 1000)
            log_api_call(request, "gigachat_meta", duration, "200 OK", 100)
            return Response({'content': res.choices[0].message.content})
        except Exception as e:
            log_api_call(request, "gigachat_meta", 0, f"Error: {str(e)[:40]}", 0)
            return Response({'content': f"Ошибка: {e}"}, status=500)

def get_task_status(request, task_id):
    task_result = AsyncResult(task_id)
    response = {'state': task_result.state, 'process': 'Обработка...', 'percent': 0}
    if task_result.state == 'PROGRESS': response.update(task_result.info)
    elif task_result.state == 'SUCCESS': response.update({'process': 'Готово!', 'percent': 100})
    elif task_result.state == 'FAILURE': response.update({'process': 'Ошибка', 'percent': 0})
    return JsonResponse(response)
