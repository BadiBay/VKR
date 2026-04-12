import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


# Сервис для сборки DOCX-документов (ТЗ для копирайтера или готовая статья)
class ExportService:

    @staticmethod
    def build_docx(
        cluster_name: str,
        doc_type: str = "tz",
        text_content: str = "",
        meta_content: str = "",
        lsi_content: str = "",
    ) -> io.BytesIO:
        """Собирает документ и возвращает его как BytesIO буфер."""
        doc = Document()

        # Базовый шрифт документа
        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)

        # Заголовок
        title_text = f"ТЗ для копирайтера: {cluster_name}" if doc_type == "tz" else f"Статья: {cluster_name}"
        head = doc.add_heading(title_text, 0)
        head.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Блок с мета-тегами
        if meta_content:
            doc.add_heading("SEO-теги (Meta)", level=1)
            doc.add_paragraph(meta_content)

        # Основной контент
        if doc_type == "article":
            doc.add_heading("Контент", level=1)
        for line in text_content.split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip())

        # LSI-слова
        if lsi_content:
            doc.add_heading("LSI слова", level=1)
            doc.add_paragraph(lsi_content)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
